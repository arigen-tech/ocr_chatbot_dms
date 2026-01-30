import os
import sqlite3
import threading
import subprocess
import pdfplumber
import pytesseract
import sys
import contextlib
import hashlib
import time
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from pathlib import Path
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from pyzbar.pyzbar import decode
from PIL import Image
import pandas as pd
import openpyxl
import json
from io import BytesIO
from Models.sql_connection.connection import execute_sql_query, execute_sql
from Models.encryption_utils import AESGCMEncryption, IllegalStateException

class DocumentHandler(FileSystemEventHandler):
    """Handler for document file system events."""
    def __init__(self, document_processor):
        self.document_processor = document_processor
        self.image_extensions = ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp']
        self.word_extensions = ['.docx', '.doc', '.docm', '.dotx', '.dotm']
        self.excel_extensions = ['.xlsx', '.xls', '.xlsm', '.csv', '.xlsb']
        self.valid_extensions = ['.pdf', '.txt'] + self.image_extensions + self.word_extensions + self.excel_extensions

    def on_created(self, event):
        file_path = event.src_path
        
        if not any(file_path.endswith(ext) for ext in self.valid_extensions):
            return
            
        print(f"\nNew document detected: {file_path}")
        # File detected - processor will handle it once DB entry is ready

class DocumentProcessor:
    _instance = None
    _lock = threading.Lock()
    
    @classmethod
    def get_instance(cls, base_dirs=["/Users/rozaltheric/Office Work/dms_project/FTP/"], db_name="search-docs.db", db_url=None, encryption_key=None):
        """Singleton accessor that supports multiple directories."""
        if not cls._instance:
            cls._instance = cls(base_dirs, db_name, db_url, encryption_key)
            cls._instance.load_existing_documents()
        return cls._instance

    def __init__(self, base_dirs, db_name, db_url, encryption_key=None):
        """Initialize with multiple base directories and optional encryption key."""
        if not base_dirs or not db_name:
            raise ValueError("base_dirs and db_name are required")
        self.base_dirs = [Path(d) for d in base_dirs]
        self.db_path = self.base_dirs[0] / db_name  # Store DB in first folder
        self.db_url = db_url
        
        # Track max ID from document_details to detect new entries
        self.max_id_lock = threading.Lock()
        self.max_id = 0
        self.polling_thread = None
        self.stop_polling = False
        
        # Initialize encryption if key provided (16 bytes for AES-128)
        self.encryption = None
        if encryption_key:
            if isinstance(encryption_key, str):
                # Convert string to bytes (ensure it's 16 bytes)
                encryption_key = encryption_key.encode('utf-8')[:16].ljust(16, b'\0')
            self.encryption = AESGCMEncryption(encryption_key)
        
        for base_dir in self.base_dirs:
            base_dir.mkdir(parents=True, exist_ok=True)
        self._initialize_db()
        
        # Initialize max_id from database
        self._update_max_id()
    
    def _get_valid_extensions(self):
        """Return a list of supported extensions."""
        return ['.pdf', '.txt', '.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp',
                '.docx', '.doc', '.docm', '.dotx', '.dotm',
                '.xlsx', '.xls', '.xlsm', '.csv', '.xlsb']

    def _initialize_db(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute("""
            CREATE VIRTUAL TABLE IF NOT EXISTS document_data USING FTS5(
                mysql_original_id UNINDEXED,
                file_name,
                content,
                qr_data,
                hash UNINDEXED
            );
        """)

        conn.commit()
        conn.close()


    def get_db_connection(self):
        """Create a new database connection for the current thread."""
        return sqlite3.connect(self.db_path)

    def _update_max_id(self):
        """Update the current max ID from document_details table (single connection, no slots wasted)."""
        if not self.db_url:
            return
        try:
            # Suppress logging for this query
            import logging
            logger = logging.getLogger('Models.sql_connection.connection')
            old_level = logger.level
            logger.setLevel(logging.WARNING)
            
            try:
                query = "SELECT MAX(id) as max_id FROM document_details"
                df = execute_sql_query(query, self.db_url)
                if not df.empty and df.iloc[0]['max_id'] is not None:
                    with self.max_id_lock:
                        self.max_id = int(df.iloc[0]['max_id'])
            finally:
                logger.setLevel(old_level)
        except Exception as e:
            print(f"Error updating max_id: {e}")

    def _get_new_entries_since_max_id(self):
        """Get newly added entries from document_details since last max_id (single connection)."""
        if not self.db_url:
            return []
        try:
            with self.max_id_lock:
                current_max = self.max_id
            
            # Suppress logging for this query since it runs frequently
            import logging
            logger = logging.getLogger('Models.sql_connection.connection')
            old_level = logger.level
            logger.setLevel(logging.WARNING)
            
            try:
                query = "SELECT id, file_name FROM document_details WHERE id > :max_id ORDER BY id ASC"
                df = execute_sql_query(query, self.db_url, params={'max_id': current_max})
                
                if not df.empty:
                    # Update max_id to the latest one we found
                    latest_id = int(df.iloc[-1]['id'])
                    with self.max_id_lock:
                        self.max_id = latest_id
                    print(f"[DB Poll] Found {len(df)} new document(s): {', '.join(df['file_name'].astype(str).tolist())}")
                    return df.to_dict('records')
                return []
            finally:
                logger.setLevel(old_level)
        except Exception as e:
            print(f"Error getting new entries: {e}")
            return []

    def _poll_for_new_documents(self):
        """Background thread that continuously polls for new document entries in DB and processes them."""
        pending_files = {}  # Track files we've seen in DB but not on disk: {file_name: timestamp}
        max_wait_time = 15  # Wait up to 15 seconds for file to appear on disk
        
        while not self.stop_polling:
            try:
                new_entries = self._get_new_entries_since_max_id()
                if new_entries:
                    for entry in new_entries:
                        file_name = entry['file_name']
                        if file_name not in pending_files:
                            pending_files[file_name] = time.time()
                            print(f"[Poll] Queued for processing: {file_name}")
                
                # Check all pending files to see if they exist on disk now
                files_to_remove = []
                for file_name, first_seen_time in list(pending_files.items()):
                    found = False
                    elapsed = time.time() - first_seen_time
                    
                    # Try to find the file recursively
                    for base_dir in self.base_dirs:
                        try:
                            for f in base_dir.rglob(file_name):
                                if f.is_file():
                                    print(f"[Poll] Found on disk after {elapsed:.1f}s: {file_name} at {f}")
                                    self.process_single_document(str(f))
                                    found = True
                                    files_to_remove.append(file_name)
                                    break
                        except Exception as search_e:
                            print(f"[Poll] Error searching for {file_name}: {search_e}")
                        
                        if found:
                            break
                    
                    # If file hasn't appeared after max_wait_time, give up
                    if not found and elapsed > max_wait_time:
                        print(f"[Poll] Timeout - file {file_name} not found after {max_wait_time}s")
                        files_to_remove.append(file_name)
                
                # Remove processed/abandoned files from pending list
                for file_name in files_to_remove:
                    del pending_files[file_name]
                
                time.sleep(1)  # Poll every second
            except Exception as e:
                print(f"[Poll] Error in polling thread: {e}")
                time.sleep(2)

    def get_id_from_db(self, file_name):
        if not self.db_url:
            return None
        query = "SELECT id FROM document_details WHERE file_name = :file_name"
        df = execute_sql_query(query, self.db_url, params={'file_name': file_name})
        if not df.empty:
            return df.iloc[0]['id']
        # If not found, do not insert, return None
        return None

    def load_existing_documents(self):
        """Scan all base directories (including subfolders) and index documents."""
        doc_files = []
        valid_exts = set(self._get_valid_extensions())

        for base_dir in self.base_dirs:
            # Use rglob to recursively find all files in nested directories
            for file_path in base_dir.rglob("*"):
                # Check if it's a file (not a directory)
                if file_path.is_file():
                    # Check if file extension matches
                    if file_path.suffix.lower() in valid_exts:
                        doc_files.append(file_path)
                        print(f"  Found: {file_path.relative_to(base_dir)}")

        if not doc_files:
            print("No Documents found across configured directories.")
            return

        print(f"\nLoading {len(doc_files)} Documents into the database...")
        for doc_path in doc_files:
            try:
                self.process_single_document(str(doc_path))
            except Exception as e:
                print(f"Error processing {doc_path.name}: {e}")

        print("Database initialization complete.")

    def extract_text_from_pdf(self, doc_path):
        """Extract text from PDF files (automatically handles both encrypted and plaintext)."""
        full_text = ""
        qr_content = set()

        # Try to decrypt if encryption is enabled
        pdf_source = doc_path
        if self.encryption:
            try:
                with open(doc_path, 'rb') as f:
                    decrypted_stream = self.encryption.decrypt_stream(f)
                pdf_source = decrypted_stream
            except (IllegalStateException, Exception):
                # File is not encrypted or decryption failed, try as plaintext
                pdf_source = doc_path

        # Helper to silence C-level stderr (zbar assertions) during decode calls
        @contextlib.contextmanager
        def _suppress_stderr():
            try:
                with open(os.devnull, 'w') as devnull:
                    old_stderr_fno = os.dup(2)
                    os.dup2(devnull.fileno(), 2)
                    try:
                        yield
                    finally:
                        os.dup2(old_stderr_fno, 2)
                        os.close(old_stderr_fno)
            except Exception:
                # If dup/dup2 not available or fails, fall back to contextlib.redirect_stderr
                with contextlib.redirect_stderr(sys.stderr):
                    yield

        try:
            with pdfplumber.open(pdf_source) as pdf:
                for page in pdf.pages:
                    try:
                        text = page.extract_text(x_tolerance=2)
                        if text:
                            full_text += text + "\n"
                    except Exception as page_text_e:
                        print(f"Warning: could not extract page text for {doc_path}: {page_text_e}")

                    # Render page to an image for OCR and QR scanning. Convert to RGB to be safe.
                    try:
                        img_pil = page.to_image(resolution=300).original.convert('RGB')
                        # img_pil = self._auto_scale_for_ocr_pil(img_pil)

                        custom_config = r'--oem 3 --psm 6 -c user_defined_dpi=300'
                        ocr_text = pytesseract.image_to_string(img_pil, config=custom_config)
                    except Exception as img_e:
                        print(f"Warning: could not render page image for OCR for {doc_path}: {img_e}")
                        continue

                    try:
                        ocr_text = pytesseract.image_to_string(img_pil)
                        if ocr_text:
                            full_text += ocr_text + "\n"
                    except Exception as ocr_e:
                        print(f"Warning: pytesseract failed on page image for {doc_path}: {ocr_e}")

                    # Decode QR codes, but suppress noisy zbar stderr and catch exceptions so one bad page
                    # won't stop the whole document processing.
                    try:
                        with _suppress_stderr():
                            qr_codes = []
                            try:
                                qr_codes = decode(img_pil)
                            except Exception as decode_e:
                                # pyzbar sometimes raises on malformed images; log and continue
                                print(f"Warning: pyzbar.decode failed for {doc_path}: {decode_e}")

                            for qr in qr_codes or []:
                                try:
                                    qr_content.add(qr.data.decode())
                                except Exception:
                                    # If decoding bytes fails, ignore this QR
                                    pass
                    except Exception as stderr_e:
                        # If suppression wrapper fails, continue without breaking whole PDF
                        print(f"Warning: failed suppressing stderr for QR decode on {doc_path}: {stderr_e}")

            # Return text (possibly empty) and qr_data (None if empty)
            return full_text.strip(), "; ".join(sorted(qr_content)) if qr_content else None

        except Exception as e:
            print(f"Error extracting from PDF {doc_path}: {str(e)}")
            return None, None
        
    def convert_to_docx(self, input_path):
        """ Convert .doc, .docm, .dotx, .dotm to .docx using unoconv. Always returns a string (output path) or None."""
        input_path = str(input_path)
        output_path = os.path.splitext(input_path)[0] + ".docx"
        try:
            subprocess.run(["unoconv", "-f", "docx", "-o", output_path, input_path], check=True)
            if os.path.exists(output_path):
                return output_path
            else:
                print(f"Conversion failed: {input_path} → {output_path}")
                return None
        except subprocess.CalledProcessError as e:
            print(f"Error converting {input_path} to .docx: {str(e)}")
            return None
    
    def extract_text_from_word(self, doc_path):
        """Extract text from various Word formats, including images (automatically handles both encrypted and plaintext)."""
        ext = os.path.splitext(doc_path)[1].lower()
        
        # Try to decrypt first if encryption is enabled
        decrypted_bytes = None
        if self.encryption:
            try:
                with open(doc_path, 'rb') as f:
                    decrypted_stream = self.encryption.decrypt_stream(f)
                    decrypted_bytes = decrypted_stream.read()
            except (IllegalStateException, Exception):
                # File is not encrypted or decryption failed, will try as plaintext
                pass
        
        # Convert to docx if needed
        if ext in ['.doc', '.docm', '.dotx', '.dotm']:
            converted_path = self.convert_to_docx(doc_path)
            if not converted_path:
                return None
            doc_path = converted_path
            ext = '.docx'
        
        try:
            # If we have decrypted bytes, use them; otherwise read from file
            if decrypted_bytes:
                doc = Document(BytesIO(decrypted_bytes))
            else:
                doc = Document(doc_path)
            
            text_data = "\n".join([para.text for para in doc.paragraphs])
            
            # Process images in the document
            image_texts = self.extract_text_from_word_images(doc)
            
            return text_data + "\n\n" + image_texts
        except Exception as e:
            print(f"Error extracting text from {doc_path}: {str(e)}")
            return None
    
    def extract_text_from_word_images(self, doc):
        """Extract text from images in the Word document."""
        image_texts = []
        instance = DocumentProcessor.get_instance()
        temp_dir = instance.base_dirs[0] / "temp_images"
        temp_dir.mkdir(parents=True, exist_ok=True)
        for rel in doc.part.rels.values():
            if rel.reltype == RT.IMAGE:
                image_part = rel.target_part
                image_bytes = image_part.blob
                image_ext = os.path.splitext(image_part.partname)[1].lower()
                image_path = temp_dir / f"image{len(image_texts)}{image_ext}"

                with open(image_path, "wb") as f:
                    f.write(image_bytes)

                try:
                    img_pil = Image.open(str(image_path)).convert('RGB')
                    # img_pil = self._auto_scale_for_ocr_pil(img_pil)

                    custom_config = r'--oem 3 --psm 6 -c user_defined_dpi=300'
                    text = pytesseract.image_to_string(img_pil, config=custom_config)
                    image_texts.append(text.strip())
                except Exception as e:
                    print(f"Error extracting text from image {image_path}: {str(e)}")

        return "\n".join(image_texts)
    
    def extract_text_from_txt(self, txt_path):
        """Extract text from .txt files (automatically handles both encrypted and plaintext)."""
        try:
            # Try to decrypt if encryption is enabled
            if self.encryption:
                try:
                    with open(txt_path, 'rb') as f:
                        decrypted_stream = self.encryption.decrypt_stream(f)
                        return decrypted_stream.read().decode('utf-8', errors='ignore')
                except (IllegalStateException, Exception):
                    # File is not encrypted or decryption failed, try as plaintext
                    pass
            
            # Read as plaintext
            with open(txt_path, 'r', encoding='utf-8') as file:
                return file.read()
        except (IllegalStateException, Exception) as e:
            print(f"Error extracting from TXT {txt_path}: {str(e)}")
            return None
    
    def extract_text_from_image(self, img_path):
        """Extract text and QR codes from images (automatically handles both encrypted and plaintext)."""
        try:
            img_pil = None
            
            # Try to decrypt if encryption is enabled
            if self.encryption:
                try:
                    with open(img_path, 'rb') as f:
                        decrypted_stream = self.encryption.decrypt_stream(f)
                        img_pil = Image.open(decrypted_stream)
                        img_pil.load()  # Verify image is valid
                        img_pil = img_pil.convert('RGB')
                except (IllegalStateException, OSError, IOError, Exception):
                    # File is not encrypted or decryption failed, try as plaintext
                    pass
            
            # If decryption failed or encryption disabled, read as plaintext
            if img_pil is None:
                try:
                    img_pil = Image.open(img_path)
                    img_pil.load()  # Verify image is valid
                    img_pil = img_pil.convert('RGB')
                except (OSError, IOError) as open_err:
                    print(f"Warning: Cannot open image {img_path}: {open_err}")
                    print(f"Image file may be corrupted or not a valid format. Skipping...")
                    return "", None
            
            text = ""
            try:
                custom_config = r'--oem 3 --psm 6 -c user_defined_dpi=300'
                text = pytesseract.image_to_string(img_pil, config=custom_config)
            except Exception as ocr_e:
                print(f"Warning: pytesseract failed on image {img_path}: {ocr_e}")

            qr_content = None
            try:
                # Suppress zbar assertions and guard decode
                with open(os.devnull, 'w') as devnull:
                    old_stderr_fno = os.dup(2)
                    os.dup2(devnull.fileno(), 2)
                    try:
                        qr_codes = decode(img_pil)
                    finally:
                        os.dup2(old_stderr_fno, 2)
                        os.close(old_stderr_fno)

                if qr_codes:
                    qr_content = "; ".join(qr.data.decode() for qr in qr_codes)
            except Exception as decode_e:
                print(f"Warning: pyzbar.decode failed for image {img_path}: {decode_e}")

            return text.strip() if text else "", qr_content
        except (IllegalStateException, Exception) as e:
            print(f"Error extracting from Image {img_path}: {str(e)}")
            return None, None
        
    def extract_text_from_excel(self, excel_path):
        """Extract text from Excel files, supporting all standard formats (automatically handles both encrypted and plaintext)."""
        try:
            excel_path_str = str(excel_path)
            
            # Try to decrypt first if encryption is enabled
            decrypted_bytes = None
            if self.encryption:
                try:
                    with open(excel_path_str, 'rb') as f:
                        decrypted_stream = self.encryption.decrypt_stream(f)
                        decrypted_bytes = decrypted_stream.read()
                except (IllegalStateException, Exception):
                    # File is not encrypted or decryption failed, will try as plaintext
                    pass
            
            if Path(excel_path_str).suffix.lower() == ".csv":
                # Handle CSV using pandas
                if decrypted_bytes:
                    from io import StringIO
                    df = pd.read_csv(StringIO(decrypted_bytes.decode('utf-8', errors='ignore')))
                else:
                    df = pd.read_csv(excel_path_str)
                return df.to_string()

            elif Path(excel_path_str).suffix.lower() in [".xls", ".xlsx", ".xlsm", ".xlsb"]:
                # Handle Excel files using openpyxl
                if decrypted_bytes:
                    wb = openpyxl.load_workbook(BytesIO(decrypted_bytes), data_only=True)
                else:
                    wb = openpyxl.load_workbook(excel_path_str, data_only=True)
                
                sheets_text = []
                for sheet in wb.worksheets:
                    for row in sheet.iter_rows(values_only=True):
                        sheets_text.append(" ".join([str(cell) for cell in row if cell]))
                
                return "\n".join(sheets_text)

            else:
                return None  # Unsupported format
        except Exception as e:
            print(f"Error extracting from Excel {excel_path}: {str(e)}")
            return None

    def is_file_in_database(self, conn, file_name):
        """Check if a file exists in the database."""
        cursor = conn.cursor()
        cursor.execute("SELECT 1 FROM document_data WHERE file_name = ?", (file_name,))
        return cursor.fetchone() is not None

    def is_hash_in_database(self, conn, hash_value):
        """Check if a hash exists in the database."""
        cursor = conn.cursor()
        cursor.execute("SELECT 1 FROM document_data WHERE hash = ?", (hash_value,))
        return cursor.fetchone() is not None
    
    def process_queue(self):
        """Process Documents from the queue continuously using parallel processing."""
        while True:
            try:
                doc_path = self.processing_queue.get()
                self.executor.submit(self.process_single_document, doc_path)  # Async processing
                self.processing_queue.task_done()
            except Exception as e:
                print(f"Error processing document from queue: {e}")
    
    def log_failed_file(self, file_name):
        """ Stores failed/corrupted/unreadable file names in failed_files.json as a JSON list. Ensures uniqueness—no duplicate names."""
        failed_log_path = self.db_path.parent / "failed_files.json"
        try:
            if failed_log_path.exists():
                with open(failed_log_path, "r", encoding="utf-8") as f:
                    existing_names = set(json.load(f))
            else:
                existing_names = set()
            
            if file_name not in existing_names:
                existing_names.add(file_name)
                with open(failed_log_path, "w", encoding="utf-8") as f:
                    json.dump(sorted(existing_names), f, ensure_ascii=False, indent=2)
        except Exception as log_e:
            print(f"[Log Error] Could not write to failed_files.json: {log_e}")

    def process_single_document(self, doc_path):
        """Process a document while checking for duplicate content."""
        conn = self.get_db_connection()
        try:
            doc_path = Path(doc_path)
            file_name = doc_path.name

            handler = DocumentHandler(self)

            if doc_path.suffix.lower() not in handler.valid_extensions:
                print(f"Skipping unsupported file type: {file_name}")
                return

            print(f"Processing: {file_name}")

            # Get id from MySQL (should be available since we waited in on_created)
            doc_id = self.get_id_from_db(file_name)
            if doc_id is None:
                print(f"File {file_name} not found in document_details, skipping")
                return

            text, qr_data = None, None
            
            try:
                # Determine file type by extension (files may be encrypted but have original extensions)
                file_ext = doc_path.suffix.lower()
                
                if file_ext in handler.image_extensions:
                    text, qr_data = self.extract_text_from_image(str(doc_path))
                elif file_ext in handler.word_extensions:
                    text = self.extract_text_from_word(str(doc_path))
                elif file_ext in handler.excel_extensions:
                    text = self.extract_text_from_excel(str(doc_path))
                elif file_ext == ".pdf":
                    text, qr_data = self.extract_text_from_pdf(str(doc_path))
                elif file_ext == ".txt":
                    text = self.extract_text_from_txt(str(doc_path))

            except Exception as inner_e:
                print(f"Exception while extracting from {file_name}: {inner_e}")
                self.log_failed_file(file_name)
                return
            
            text_is_empty = (text is None or (isinstance(text, str) and not text.strip()))

            if text_is_empty:
                print(f"Failed to extract text from: {file_name}")
                self.log_failed_file(file_name)
                return

            # Compute hash based on content ONLY (not filename, to detect duplicates across versions)
            normalized_text = " ".join(text.split())

            # Create fingerprint without filename for duplicate detection
            content_fingerprint = (
                normalized_text[:5000] + "|" +
                (qr_data or "")
            )

            hash_value = hashlib.sha256(content_fingerprint.encode("utf-8")).hexdigest()

            # Check if hash exists in search-docs.db (duplicate content detection)
            cursor = conn.cursor()
            cursor.execute(
                "SELECT mysql_original_id FROM document_data WHERE hash = ? ORDER BY mysql_original_id ASC LIMIT 1",
                (hash_value,)
            )
            row = cursor.fetchone()
            
            if row is not None and row[0] is not None:
                # Hash found with valid original ID
                found_original_id = int(row[0])
                current_id = int(doc_id)
                
                # Determine which is truly original: the one with LOWER ID
                if current_id < found_original_id:
                    # Current file has lower ID → current is ORIGINAL, found one is DUPLICATE
                    # Update the FTS record to point to the true original
                    cursor.execute(
                        """
                        UPDATE document_data
                        SET mysql_original_id = ?
                        WHERE mysql_original_id = ?
                        """,
                        (current_id, found_original_id)
                    )
                    conn.commit()
                    
                    # Mark the found_original_id as a duplicate of current
                    execute_sql(
                        """
                        UPDATE document_details
                        SET is_duplicate = 1,
                            document_id = :original_id
                        WHERE id = :duplicate_id
                        """,
                        self.db_url,
                        {
                            'original_id': current_id,
                            'duplicate_id': found_original_id
                        }
                    )
                    
                    # Mark current as original
                    execute_sql(
                        """
                        UPDATE document_details
                        SET is_duplicate = 0,
                            document_id = NULL
                        WHERE id = :id
                        """,
                        self.db_url,
                        {'id': current_id}
                    )
                    print(f"[Original] {file_name} (id={current_id}) is original, marking {found_original_id} as duplicate")
                    return
                else:
                    # Current file has higher ID → current is DUPLICATE, found one is ORIGINAL
                    execute_sql(
                        """
                        UPDATE document_details
                        SET is_duplicate = 1,
                            document_id = :original_id
                        WHERE id = :current_id
                        """,
                        self.db_url,
                        {
                            'original_id': found_original_id,
                            'current_id': current_id
                        }
                    )
                    print(f"[Duplicate] {file_name} → original id {found_original_id}")
                    return
            
            elif row is not None and row[0] is None:
                # Corrupted FTS row (NULL original) - fix it
                print(f"[REPAIR] Fixing NULL original for hash {hash_value}")
                cursor.execute(
                    """
                    UPDATE document_data
                    SET mysql_original_id = ?
                    WHERE hash = ?
                    """,
                    (int(doc_id), hash_value)
                )
                conn.commit()
                
                # Mark as original since we fixed it
                execute_sql(
                    """
                    UPDATE document_details
                    SET is_duplicate = 0,
                        document_id = NULL
                    WHERE id = :id
                    """,
                    self.db_url,
                    {'id': int(doc_id)}
                )
                return
            
            # Content hash NOT FOUND - this is a NEW/ORIGINAL file
            # Insert its content into search-docs.db
            cursor.execute(
                """
                INSERT INTO document_data
                (mysql_original_id, file_name, content, qr_data, hash)
                VALUES (?, ?, ?, ?, ?)
                """,
                (int(doc_id), file_name, text, qr_data, hash_value)
            )
            conn.commit()
            print(f"[Original] Inserted {file_name} (id={doc_id}) as original file")
            
            # Update MySQL: Mark current file as original (not a duplicate)
            execute_sql(
                """
                UPDATE document_details
                SET is_duplicate = 0,
                    document_id = NULL
                WHERE id = :id
                """,
                self.db_url,
                {'id': int(doc_id)}
            )
                
        finally:
            conn.close()

    def get_all_documents(self):
        """Retrieve all Documents names stored in the database."""
        conn = self.get_db_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT mysql_original_id, file_name FROM document_data")
            return [
                {
                    "mysql_original_id": row[0],
                    "file_name": row[1]
                } for row in cursor.fetchall()]
        finally:
            conn.close()

    def search_database(self, query, selected_files=None):
        """Search the database and return mysql_original_id + file_name."""
        conn = self.get_db_connection()
        try:
            cursor = conn.cursor()

            fts_query = f'"{query}"'
            wildcard_query = f"%{query}%"

            placeholders = None
            if selected_files:
                placeholders = ','.join('?' * len(selected_files))

            if selected_files:
                match_sql = f"""
                    SELECT DISTINCT mysql_original_id, file_name
                    FROM document_data
                    WHERE content MATCH ?
                    AND mysql_original_id IN ({placeholders})
                """
                cursor.execute(match_sql, (fts_query, *selected_files))
            else:
                match_sql = """
                    SELECT DISTINCT mysql_original_id, file_name
                    FROM document_data
                    WHERE content MATCH ?
                """
                cursor.execute(match_sql, (fts_query,))

            rows = cursor.fetchall()

            if not rows:
                if selected_files:
                    like_sql = f"""
                        SELECT DISTINCT mysql_original_id, file_name
                        FROM document_data
                        WHERE content LIKE ?
                        AND mysql_original_id IN ({placeholders})
                    """
                    cursor.execute(like_sql, (wildcard_query, *selected_files))
                else:
                    like_sql = """
                        SELECT DISTINCT mysql_original_id, file_name
                        FROM document_data
                        WHERE content LIKE ?
                    """
                    cursor.execute(like_sql, (wildcard_query,))

                rows = cursor.fetchall()

            results = [
                {
                    "mysql_original_id": row[0],
                    "file_name": row[1]
                }
                for row in rows
            ]

            return results

        finally:
            conn.close()
    
    def clean_database(self):
        """Delete all entries in the document_data table, effectively resetting the database."""
        conn = self.get_db_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM document_data")
            conn.commit()
            print("Database cleaned successfully.")
        except sqlite3.Error as e:
            print(f"Database error while cleaning: {str(e)}")
        finally:
            conn.close()

    

    def start_processing(self):
        """Begin monitoring all folders with Watchdog and start polling thread."""
        # self.load_existing_documents()

        # Start the polling thread for new database entries
        self.stop_polling = False
        self.polling_thread = threading.Thread(target=self._poll_for_new_documents, daemon=True)
        self.polling_thread.start()
        print("Document polling thread started")

        # Also start Watchdog for real-time file detection
        observer = Observer()
        event_handler = DocumentHandler(self)

        for base_dir in self.base_dirs:
            observer.schedule(event_handler, str(base_dir), recursive=True)

        observer.start()
        return observer