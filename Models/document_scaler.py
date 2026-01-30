"""
Standalone Document Scaling Service
This module provides document scaling functionality for PDF, Images, and Word documents.
Can be integrated with Spring Boot or other applications via API or direct import.
"""

import os
from pathlib import Path
from PIL import Image
from io import BytesIO
import numpy as np
import cv2

# Import PDF handling libraries
try:
    from pdf2image import convert_from_path
    POPPLER_AVAILABLE = True
except ImportError:
    POPPLER_AVAILABLE = False

try:
    from pypdf import PdfReader, PdfWriter
except ImportError:
    PdfReader = None
    PdfWriter = None

# Import Word document handling
try:
    from docx import Document
except ImportError:
    Document = None


class DocumentScaler:
    """
    Standalone document scaling service for multiple file formats.
    Supports upscaling (150%) and downscaling (66%).
    """
    
    # Supported file formats
    IMAGE_FORMATS = ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp']
    WORD_FORMATS = ['.docx', '.doc', '.docm', '.dotx', '.dotm']
    PDF_FORMATS = ['.pdf']
    SUPPORTED_FORMATS = PDF_FORMATS + IMAGE_FORMATS + WORD_FORMATS
    
    def __init__(self, output_base_dir=None, base_path=None):
        """
        Initialize the document scaler.
        
        Args:
            output_base_dir (str, optional): Base directory for scaled documents.
            Defaults to a 'scaled_documents' folder
            in the same directory as the input file.
            base_path (str, optional): Base directory to search for files if file_path is not absolute.
        """
        # Set default base_path to current working directory if not provided
        if base_path is None:
            base_path = os.getcwd()
        self.base_path = Path(base_path)
        
        self.output_base_dir = output_base_dir
        self._validate_dependencies()
    
    def _validate_dependencies(self):
        """Validate that required libraries are installed."""
        if Document is None:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    def _find_file(self, file_path, base_path):
        """
        Find the file in the base_path and subfolders if file_path is not absolute.
        
        Args:
            file_path (str): The file path or name
            base_path (Path or None): Base directory to search
        
        Returns:
            Path or None: The found file path
        """
        file_path_obj = Path(file_path)
        
        # If absolute path and exists, return it
        if file_path_obj.is_absolute() and file_path_obj.exists():
            return file_path_obj
        
        # If base_path is provided, search there
        if base_path and base_path.exists():
            # If file_path is absolute but doesn't exist, search for basename
            search_name = file_path_obj.name if file_path_obj.is_absolute() else file_path
            
            # First, check if it's a direct path relative to base
            potential_path = base_path / search_name
            if potential_path.exists():
                return potential_path
            
            # Then, search recursively for the name
            for found_path in base_path.rglob(search_name):
                return found_path
        
        # If no base_path or not found, check if relative to current dir exists
        if file_path_obj.exists():
            return file_path_obj
        
        return None
    
    def scale_file(self, file_path, scale_type, output_dir=None, base_path=None):
        """ Scale a document file. """
        # Validate inputs
        if scale_type not in [0, 1]:
            raise ValueError("scale_type must be 0 (downscale) or 1 (upscale)")
        
        # Determine base_path for search
        search_base = Path(base_path) if base_path else self.base_path
        
        # Find the actual file path
        file_path_obj = self._find_file(file_path, search_base)
        if not file_path_obj:
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path_obj.suffix.lower()
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported: {', '.join(self.SUPPORTED_FORMATS)}")
        
        # Determine scale factor and label
        scale_factor = 1.5 if scale_type == 1 else 0.66
        scale_label = "upscaled" if scale_type == 1 else "downscaled"
        
        # For overwrite: output path is the same as input path
        output_path = file_path_obj
        
        try:
            # Route to appropriate scaling method
            if file_ext == ".pdf":
                self._scale_pdf(str(file_path_obj), str(output_path), scale_factor)
            elif file_ext in self.IMAGE_FORMATS:
                self._scale_image(str(file_path_obj), str(output_path), scale_factor)
            elif file_ext in self.WORD_FORMATS:
                self._scale_word_document(str(file_path_obj), str(output_path), scale_factor)
            
            print(f"✓ Document scaled successfully: {output_path}")
            return str(output_path)
        
        except Exception as e:
            print(f"✗ Error scaling document: {e}")
            raise
    
    def _scale_pdf(self, input_path, output_path, scale_factor):
        """
        Scale PDF by converting to images, scaling them, and converting back.
        Falls back to PyPDF2 if poppler is not available.
        """
        if not POPPLER_AVAILABLE:
            print("⚠ Poppler not available, using PyPDF2 fallback...")
            return self._scale_pdf_pypdf2(input_path, output_path, scale_factor)
        
        try:
            # Adjust DPI based on scale factor to control file size
            base_dpi = 150 if scale_factor < 1 else 200
            dpi = int(base_dpi * scale_factor)
            
            # Convert PDF to images
            images = convert_from_path(input_path, dpi=dpi)
            
            if not images:
                raise ValueError("No pages found in PDF")
            
            # Scale images
            scaled_images = []
            for img in images:
                new_width = int(img.width * scale_factor)
                new_height = int(img.height * scale_factor)
                scaled_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                scaled_images.append(scaled_img)
            
            # Save back to PDF with optimization
            scaled_images[0].save(
                output_path,
                save_all=True,
                append_images=scaled_images[1:] if len(scaled_images) > 1 else [],
                format='PDF',
                optimize=True,
                quality=85 if scale_factor < 1 else 95
            )
            print(f"  → PDF scaled using poppler")
        
        except Exception as pdf_err:
            print(f"  → Poppler error: {pdf_err}")
            print(f"  → Attempting PyPDF2 fallback...")
            self._scale_pdf_pypdf2(input_path, output_path, scale_factor)
    
    def _scale_pdf_pypdf2(self, input_path, output_path, scale_factor):
        """
        Fallback PDF scaling using PyPDF2 (pure Python, no external dependencies).
        """
        if PdfReader is None or PdfWriter is None:
            raise ImportError(
                "PyPDF2 not installed. Install with: pip install pypdf\n"
                "Or use poppler for PDF scaling: brew install poppler (macOS)"
            )
        
        try:
            reader = PdfReader(input_path)
            writer = PdfWriter()
            
            for page in reader.pages:
                page.scale(scale_factor, scale_factor)
                writer.add_page(page)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            print(f"  → PDF scaled using PyPDF2")
        
        except Exception as e:
            print(f"✗ PDF scaling failed with both methods: {e}")
            raise
    
    def _scale_image(self, input_path, output_path, scale_factor):
        """Scale image file with robust error handling for corrupted images."""
        try:
            try:
                img = Image.open(input_path)
                img.load()  # Verify image is valid
            except (OSError, IOError) as img_err:
                raise ValueError(f"Invalid or corrupted image file: {input_path}. Error: {img_err}")
            
            new_width = int(img.width * scale_factor)
            new_height = int(img.height * scale_factor)
            
            scaled_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # Adjust quality based on scale factor
            quality = 85 if scale_factor < 1 else 95
            scaled_img.save(output_path, quality=quality)
            
            print(f"  → Image scaled ({int(scale_factor * 100)}%)")
        
        except Exception as e:
            print(f"✗ Image scaling failed: {e}")
            raise
    
    def _scale_word_document(self, input_path, output_path, scale_factor):
        """Scale images in Word document and create new document."""
        try:
            doc = Document(input_path)
            new_doc = Document()
            
            # Copy paragraphs with text formatting
            for para in doc.paragraphs:
                new_para = new_doc.add_paragraph(para.text)
                new_para.style = para.style
                
                # Copy runs with formatting
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.font.size = run.font.size
                    new_run.font.bold = run.font.bold
                    new_run.font.italic = run.font.italic
            
            # Process and scale images
            image_count = 0
            for rel in doc.part.rels.values():
                if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image':
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    
                    try:
                        # Scale the image
                        img = Image.open(BytesIO(image_bytes))
                        new_width = int(img.width * scale_factor)
                        new_height = int(img.height * scale_factor)
                        scaled_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                        
                        # Save scaled image
                        img_byte_arr = BytesIO()
                        scaled_img.save(img_byte_arr, format=image_part.partname.split('.')[-1].upper())
                        img_byte_arr.seek(0)
                        
                        # Add to new document
                        new_doc.add_picture(img_byte_arr, width=int(scaled_img.width * 914400 / 96))
                        image_count += 1
                    except Exception as img_err:
                        print(f"  ⚠ Warning: Could not scale embedded image: {img_err}")
            
            new_doc.save(output_path)
            print(f"  → Word document scaled ({image_count} images processed)")
        
        except Exception as e:
            print(f"✗ Word document scaling failed: {e}")
            raise
    
    def auto_scale_for_ocr(self, img_pil, min_width=1500, max_width=3000):
        """
        Adaptive scaling for OCR optimization (in-memory).
        Does NOT affect original or scaled files.
        
        Args:
            img_pil (PIL.Image): PIL Image object
            min_width (int): Minimum width threshold
            max_width (int): Maximum width threshold
        
        Returns:
            PIL.Image: Scaled image or original if scaling not needed
        """
        try:
            img_np = np.array(img_pil)
            h, w = img_np.shape[:2]
            
            scale = 1.0
            
            if w < min_width:
                scale = min_width / w
            elif w > max_width:
                scale = max_width / w
            
            # Skip if scaling is insignificant
            if abs(scale - 1.0) < 0.05:
                return img_pil
            
            new_w = int(w * scale)
            new_h = int(h * scale)
            
            resized = cv2.resize(img_np, (new_w, new_h), interpolation=cv2.INTER_CUBIC)
            return Image.fromarray(resized)
        
        except Exception as e:
            print(f"[OCR Scale] Warning: auto scale failed: {e}")
            return img_pil
    
    @staticmethod
    def get_supported_formats():
        """Get list of supported file formats."""
        return DocumentScaler.SUPPORTED_FORMATS
    
    @staticmethod
    def validate_file(file_path):
        """
        Validate if a file is supported.
        
        Args:
            file_path (str): Path to file
        
        Returns:
            bool: True if supported, False otherwise
        """
        file_ext = Path(file_path).suffix.lower()
        return file_ext in DocumentScaler.SUPPORTED_FORMATS


# Convenience functions for direct usage
def scale_document(file_path, scale_type, output_dir=None, base_path=None):
    """
    Convenience function to scale a document.
    
    Args:
        file_path (str): Path to the file
        scale_type (int): 1 for upscale, 0 for downscale
        output_dir (str, optional): Custom output directory
        base_path (str, optional): Base directory to search for files (defaults to current dir)
    
    Returns:
        str: Path to scaled file
    """
    scaler = DocumentScaler(base_path=base_path)
    return scaler.scale_file(file_path, scale_type, output_dir, base_path)


def auto_scale_for_ocr(img_pil, min_width=1500, max_width=3000):
    """
    Convenience function for OCR adaptive scaling.
    
    Args:
        img_pil (PIL.Image): PIL Image object
        min_width (int): Minimum width threshold
        max_width (int): Maximum width threshold
    
    Returns:
        PIL.Image: Scaled image
    """
    scaler = DocumentScaler()
    return scaler.auto_scale_for_ocr(img_pil, min_width, max_width)


if __name__ == "__main__":
    # Example usage
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python document_scaler.py <file_path> <scale_type> [base_path] [output_dir]")
        print("  scale_type: 1 for upscale (150%), 0 for downscale (66%)")
        print("  base_path: Base directory to search for files (optional)")
        sys.exit(1)
    
    file_path = sys.argv[1]
    scale_type = int(sys.argv[2])
    base_path = sys.argv[3] if len(sys.argv) > 3 else None
    output_dir = sys.argv[4] if len(sys.argv) > 4 else None
    
    # Set base path here
    base_path = "/Users/rozaltheric/Office Work/ocr_chatbot_dms/PDFs_Data/"  # Replace with your desired path
    output_dir = base_path
    
    try:
        result = scale_document(file_path, scale_type, output_dir, base_path)
        print(f"\nScaled file: {result}")
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
